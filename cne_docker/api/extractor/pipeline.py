import os, re
from typing import Optional, Dict, List
import pandas as pd
from docx import Document

from .rules import normalize_whitespace, split_candidates_by_type
from .ai import normalize_sigla, guess_is_name

CSV_COLUMNS = ["DTMNFR","ORGAO","TIPO","SIGLA","SIMBOLO","NOME_LISTA","NUM_ORDEM","NOME_CANDIDATO","PARTIDO_PROPONENTE","INDEPENDENTE"]

HEADER_RE = re.compile(r"^\s*([A-Z0-9\./\-]{2,12})\s*[-–—]\s*(.+?)\s*$")
ORG_SECTION_RE = re.compile(r"^\s*(\d+\s*[.)-]\s*)?(Assembleia Municipal|C[aâ]mara Municipal)\b", re.IGNORECASE)

def parse_docx(path: str) -> str:
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            lines.append(txt)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        lines.append(txt)
    return normalize_whitespace("\n".join(lines))

def is_sigla_like(token: str) -> bool:
    return bool(token) and token.upper() == token and len(token) <= 12

def infer_dtmnfr_from_path(path: str) -> str:
    m = re.search(r'_(\d{4})_', os.path.basename(path))
    if not m:
        return "0000000000"
    return f"{m.group(1)}000000"

def extract_blocks_with_orgao(text: str):
    blocks = []
    orgao_ctx = ""
    cur = None
    def push():
        nonlocal cur
        if cur:
            cur["content"] = "\n".join(cur["content"]).strip()
            blocks.append(cur); cur = None
    for raw in text.splitlines():
        l = raw.strip()
        if not l:
            continue
        m_org = ORG_SECTION_RE.match(l)
        if m_org:
            word = m_org.group(2).lower()
            if "assembleia" in word: orgao_ctx = "AM"
            elif "câmara" in word or "camara" in word: orgao_ctx = "CM"
            else: orgao_ctx = ""
            continue
        m = HEADER_RE.match(l)
        if m and is_sigla_like(m.group(1)):
            push()
            cur = {"orgao": orgao_ctx, "header": l, "content": []}
            continue
        if cur:
            cur["content"].append(l)
    push()
    return blocks

def parse_header(header_line: str, enable_ia: bool):
    m = HEADER_RE.match(header_line)
    if m:
        sigla, nome = m.group(1), m.group(2)
    else:
        parts = header_line.split(maxsplit=1)
        sigla = parts[0]; nome = parts[1] if len(parts)>1 else parts[0]
    sigla = normalize_sigla(sigla.strip()) if enable_ia else sigla.strip()
    return sigla, nome.strip()

def to_rows_from_block(header_line: str, content: str, orgao_hint: str, dtmnfr: str, ord_reset: bool, enable_ia: bool) -> List[dict]:
    sigla, nome_lista = parse_header(header_line, enable_ia)
    orgao = orgao_hint or ""
    efetivos, suplentes = split_candidates_by_type(content)
    efetivos = [c for c in efetivos if guess_is_name(c, enable_ia=enable_ia)]
    suplentes = [c for c in suplentes if guess_is_name(c, enable_ia=enable_ia)]
    rows: List[dict] = []
    def build_row(tipo, num, cand):
        return {
            "DTMNFR": dtmnfr,
            "ORGAO": orgao,
            "TIPO": tipo,
            "SIGLA": sigla,
            "SIMBOLO": "",
            "NOME_LISTA": nome_lista,
            "NUM_ORDEM": num,
            "NOME_CANDIDATO": cand.strip(),
            "PARTIDO_PROPONENTE": sigla,
            "INDEPENDENTE": "N",
        }
    if ord_reset:
        for i, c in enumerate(efetivos, start=1):
            rows.append(build_row(2, i, c))
        for j, c in enumerate(suplentes, start=1):
            rows.append(build_row(3, j, c))
    else:
        n = 1
        for c in efetivos:
            rows.append(build_row(2, n, c)); n += 1
        for c in suplentes:
            rows.append(build_row(3, n, c)); n += 1
    return rows

def extract_to_csv(in_path: str, out_csv: str, orgao: Optional[str]=None, ord_reset: bool=True, enable_ia: bool=True, models_dir: Optional[str]=None) -> Dict:
    dtmnfr = infer_dtmnfr_from_path(in_path)
    text = parse_docx(in_path)
    blocks = extract_blocks_with_orgao(text)
    all_rows: List[dict] = []
    for b in blocks:
        all_rows.extend(to_rows_from_block(
            header_line=b["header"],
            content=b["content"],
            orgao_hint=orgao or b["orgao"],
            dtmnfr=dtmnfr,
            ord_reset=ord_reset,
            enable_ia=enable_ia
        ))
    if not all_rows:
        for b in blocks:
            all_rows.extend(to_rows_from_block(
                header_line=b["header"],
                content=b["content"],
                orgao_hint=orgao or b["orgao"],
                dtmnfr=dtmnfr,
                ord_reset=ord_reset,
                enable_ia=False
            ))
    df = pd.DataFrame(all_rows, columns=CSV_COLUMNS).fillna("")
    if not df.empty:
        df = df.sort_values(by=["ORGAO","SIGLA","TIPO","NOME_LISTA","NUM_ORDEM","NOME_CANDIDATO"]).reset_index(drop=True)
    df.to_csv(out_csv, index=False, sep=";", encoding="utf-8")
    return {
        "rows": int(df.shape[0]),
        "orgoes": sorted(list(set(df["ORGAO"].unique()))) if "ORGAO" in df else [],
        "siglas": sorted(list(set(df["SIGLA"].unique()))) if "SIGLA" in df else [],
    }
